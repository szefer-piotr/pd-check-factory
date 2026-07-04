#!/usr/bin/env python3
"""Convert PD Check Factory Markdown documentation to Word format."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass(frozen=True)
class TitlePage:
    title: str
    subtitle: str
    version: str = "1.0 draft"
    status: str = "Draft for review"
    date: str = "23 June 2026"
    product: str = "Rho PD Assurance"
    classification: str = "Internal use"


def _insert_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run.add_text("Table of contents — open in Microsoft Word and update fields.")
    run._r.append(fld_char_end)


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_title_page(doc: Document, page: TitlePage) -> None:
    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run(page.title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run(page.subtitle)
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("Document version", page.version),
        ("Status", page.status),
        ("Date", page.date),
        ("Product", page.product),
        ("Classification", page.classification),
    ]
    for i, (label, value) in enumerate(rows):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = value
        _set_cell_shading(meta.rows[i].cells[0], "E8EEF4")

    doc.add_paragraph()
    doc.add_page_break()

    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    _insert_toc(doc.add_paragraph())
    doc.add_page_break()


def _parse_table_row(line: str) -> list[str] | None:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return None
    return [cell.strip() for cell in stripped.strip("|").split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-+:?", cell.replace(" ", "")) or cell == "" for cell in cells)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text
            if r_idx == 0:
                _set_cell_shading(cell, "E8EEF4")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_list_item(doc: Document, text: str, *, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    doc.add_paragraph(text, style=style)


def _convert_inline_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def convert_markdown_to_docx(
    md_path: Path,
    docx_path: Path,
    *,
    title_page: TitlePage,
) -> None:
    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    _add_title_page(doc, title_page)

    i = 0
    in_code_block = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if i == 0 and line.startswith("# "):
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.strip().startswith("```"):
            if in_code_block:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        table_cells = _parse_table_row(line)
        if table_cells is not None:
            if not _is_separator_row(table_cells):
                table_rows.append(table_cells)
            i += 1
            continue
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text = heading_match.group(2).strip()
            heading = doc.add_heading(text, level=level)
            if heading.runs:
                heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            i += 1
            continue

        ordered_match = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if ordered_match:
            _add_list_item(doc, ordered_match.group(2), ordered=True)
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", line.strip())
        if bullet_match:
            _add_list_item(doc, bullet_match.group(1), ordered=False)
            i += 1
            continue

        if line.strip() == "*End of document*":
            i += 1
            continue

        if line.strip():
            paragraph = doc.add_paragraph()
            _convert_inline_bold(paragraph, line.strip())
        i += 1

    if table_rows:
        _add_table(doc, table_rows)
    if code_lines:
        _add_code_block(doc, code_lines)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = 1
    run = footer.add_run(f"Generated from {md_path.as_posix()}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
